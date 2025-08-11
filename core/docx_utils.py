import os
from pathlib import Path
from typing import List, Dict, Any, Tuple
import logging
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import mammoth
import docx2python

logger = logging.getLogger(__name__)

class DocxProcessor:
    def __init__(self):
        self.temp_dir = None
    
    def extract_text(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text and metadata from DOCX file"""
        text = ""
        metadata = {
            'title': None,
            'author': None,
            'created': None,
            'paragraphs': 0,
            'tables': 0
        }
        
        try:
            # Primary method: python-docx
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            text = '\n'.join(paragraphs)
            metadata['paragraphs'] = len(paragraphs)
            metadata['tables'] = len(doc.tables)
            
            # Extract core properties
            if doc.core_properties:
                metadata['title'] = doc.core_properties.title
                metadata['author'] = doc.core_properties.author
                metadata['created'] = doc.core_properties.created
            
            # Fallback if no text extracted
            if not text.strip():
                raise ValueError("No text extracted with python-docx")
                
        except Exception as e:
            logger.warning(f"python-docx extraction failed: {e}, trying fallback methods")
            
            try:
                # Fallback method 1: mammoth
                with open(file_path, 'rb') as docx_file:
                    result = mammoth.extract_raw_text(docx_file)
                    text = result.value
                    
                if not text.strip():
                    raise ValueError("No text extracted with mammoth")
                    
            except Exception as e2:
                logger.warning(f"mammoth extraction failed: {e2}, trying docx2python")
                
                try:
                    # Fallback method 2: docx2python
                    content = docx2python.docx2python(file_path)
                    text = '\n'.join(content.text)
                    
                    if not text.strip():
                        raise ValueError("No text extracted with docx2python")
                        
                except Exception as e3:
                    logger.error(f"All extraction methods failed: {e3}")
                    raise ValueError("Could not extract text from document")
        
        return text, metadata
    
    def add_comments(self, file_path: Path, issues: List[Dict]) -> Path:
        """Add comments to DOCX file for identified issues"""
        if not issues:
            return file_path
        
        try:
            doc = Document(file_path)
            
            # Group issues by severity for color coding
            severity_colors = {
                'High': RGBColor(255, 0, 0),    # Red
                'Medium': RGBColor(255, 165, 0), # Orange  
                'Low': RGBColor(255, 255, 0)     # Yellow
            }
            
            for issue in issues:
                comment_text = self._format_comment(issue)
                
                # Find appropriate location to add comment
                target_para = self._find_target_paragraph(doc, issue)
                
                if target_para:
                    # Highlight the paragraph
                    severity = issue.get('severity', 'Medium')
                    color = severity_colors.get(severity, RGBColor(255, 165, 0))
                    
                    for run in target_para.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    
                    # Add comment (simplified - Word comments require more complex implementation)
                    # For now, add as a new paragraph with the comment
                    comment_para = doc.add_paragraph()
                    comment_run = comment_para.add_run(f"[COMMENT] {comment_text}")
                    comment_run.font.color.rgb = color
                    comment_run.italic = True
            
            # Save as new file
            output_path = file_path.parent / f"reviewed_{file_path.name}"
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to add comments to {file_path}: {e}")
            return file_path
    
    def _format_comment(self, issue: Dict) -> str:
        """Format issue as comment text"""
        comment_parts = [
            f"ISSUE: {issue.get('issue', 'Unknown issue')}",
            f"SEVERITY: {issue.get('severity', 'Medium')}"
        ]
        
        if issue.get('citations'):
            citations = ', '.join(issue['citations'])
            comment_parts.append(f"REFERENCE: {citations}")
        
        if issue.get('suggestion'):
            comment_parts.append(f"SUGGESTION: {issue['suggestion']}")
        
        return ' | '.join(comment_parts)
    
    def _find_target_paragraph(self, doc: Document, issue: Dict) -> Any:
        """Find the best paragraph to attach comment to"""
        
        section_hint = issue.get('section_hint', '').lower()
        issue_text = issue.get('issue', '').lower()
        
        # Search for relevant keywords in paragraphs
        search_terms = []
        
        if 'jurisdiction' in section_hint:
            search_terms = ['jurisdiction', 'governing law', 'courts', 'disputes']
        elif 'registered office' in section_hint:
            search_terms = ['registered office', 'address', 'adgm']
        elif 'signature' in section_hint:
            search_terms = ['signature', 'signed', 'witness', 'executed']
        elif 'ubo' in section_hint or 'beneficial' in section_hint:
            search_terms = ['beneficial owner', 'ubo', 'ownership', 'control']
        
        # Find best matching paragraph
        best_match = None
        best_score = 0
        
        for para in doc.paragraphs:
            para_text = para.text.lower()
            score = sum(1 for term in search_terms if term in para_text)
            
            if score > best_score:
                best_score = score
                best_match = para
        
        # If no specific match, try to find by document structure
        if not best_match:
            for para in doc.paragraphs:
                if len(para.text) > 50:  # Substantial paragraph
                    best_match = para
                    break
        
        return best_match
